import pandas as pd
from pandas import DataFrame
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Pt
from docx.oxml.ns import qn
from docx import Document
from PIL import Image
from io import BytesIO
from pathlib import Path
from datetime import datetime
import time
from src.LOG_DATA_STEEL import LOG_DICT
from src.interraction_terminal import set_argumments

def clean_and_save(doc_path, keyword="待删除段落"):
    """一轮清理：删除包含关键字的段落"""
    doc = Document(doc_path)
    
    for para in reversed(doc.paragraphs):
        # 删除条件：包含关键字 
        if keyword in para.text :
            para._element.getparent().remove(para._element)

    if doc.paragraphs:
        for br in reversed(list(doc.paragraphs[-1]._element.iter(qn('w:br')))):
            if br.get(qn('w:type')) == 'page':
                br.getparent().remove(br)
                break
    
    doc.save(doc_path)

def compress_image(path, max_width_px=1200, quality=85):
    """
    压缩图片并返回 BytesIO，供 InlineImage 直接使用
    """
    try:
        with Image.open(path) as img:
            # 如果尺寸过大，缩小
            if img.width > max_width_px:
                ratio = max_width_px / img.width
                new_size = (max_width_px, int(img.height * ratio))
                img = img.resize(new_size, Image.LANCZOS)
            
            # 压缩到内存 buffer
            buffer = BytesIO()
            if img.mode == 'RGBA':
                img.save(buffer, format='PNG')
            else:
                img.save(buffer, format='JPEG', quality=quality, optimize=True)
            buffer.seek(0)
            return buffer
    except Exception as e:
        print(f"压缩失败 {path}: {e}")
        return path  # 失败时返回原路径，让 InlineImage 自己处理

def make_data_in_list(tpl,df:DataFrame=None,dfs:dict[str,DataFrame]=None,key='')->dict:
    """编制填表的内容"""
    data_dict={}
    #   封面
    df=dfs['封面']
    df=df[df['报告编号']==key]
    df=df.add_prefix('封面_')
    data_dict.update(df.iloc[0].to_dict())
    #   评估情况表
    df=dfs['评估情况表']
    df=df[df['报告编号']==key]
    for word in ['立即改造','符合安全运行要求','限期改造','落实安全管控措施，可继续运行']:
        if word=='落实安全管控措施，可继续运行':
            df['落实安全管控措施']=f'☑{word}' if  word in df['评估结果'].iloc[0] else f'□{word}'
        else:
            df[f'{word}']=f'☑{word}' if  word in df['评估结果'].iloc[0] else f'□{word}'
    ps:str=df['主要问题'].iloc[0]
    for word in ['材质落后','使用年限较长','腐蚀泄漏严重','防腐状况较差','建构筑物占压','处于或临近地质灾害易发区域','处于或临近人员密集区']:
        df[f'{word}']=f'☑{word}' if  word in ps.split('：')[0] else f'□{word}'
    if len(ps.split('：'))>1  and ps.split('：')[1].strip() != '':
        df['其他主要问题']=f'☑其他主要问题：{ps.split('：')[1].strip()}'
    else:
        df['其他主要问题']='□其他主要问题：'
    df=df.add_prefix('评估情况表_')
    data_dict.update(df.iloc[0].to_dict())
    #   资料审查
    df=dfs['资料审查']
    df=df[df['报告编号']==key]
    large=len(df)
    df=df.fillna('不明')
    df['竣工验收日期']=df['竣工验收日期'].apply(lambda x: x.strftime('%Y年%m月%d日') if isinstance(x, pd.Timestamp) else x)
    data_dict.update({'资料审查_记数':large})
    text = ';'.join(df['资料审查问题记载'].dropna().unique().astype(str))
    data_dict.update({'资料审查_资料审查总结':text})
    temp_dict={'资料审查报告':df.to_dict(orient='records')}
    data_dict.update(temp_dict)
    #   庭院钢管宏观检查
    df=dfs['庭院钢管宏观检查']
    df=df[df['报告编号']==key]
    text = ';'.join(df['结论'].dropna().unique().astype(str))
    data_dict.update({'庭院钢管_总结':text})
    temp_dict={'庭院钢管检查报告':df.to_dict(orient='records')}
    data_dict.update(temp_dict)
    #   立管宏观检查
    # df=dfs['立管宏观检查']
    # df=df[df['报告编号']==key]
    # text = ';'.join(df['结论'].dropna().unique().astype(str))
    # data_dict.update({'立管_总结':text})
    # temp_dict={'立管宏观检查':df.to_dict(orient='records')}
    # data_dict.update(temp_dict)
    #   泄漏检测
    df=dfs['泄漏检测']
    df=df[df['报告编号']==key]
    df['检测时间']=df['检测时间'].apply(lambda x: x.strftime('%Y年%m月') if isinstance(x, pd.Timestamp) else x)
    text = ';'.join(df['检测结果'].dropna().unique().astype(str))
    data_dict.update({'泄漏_总结':text})
    group_cols = ['管道名称', '管道材质', '管道位置','检测结果']
    value_cols = ['检测点位置', '检测时间', '浓度']
    
    df = df[group_cols + value_cols].copy()

    # 2. 逐组补齐到7行
    dfz = []
    for name, g in df.groupby(group_cols):
        n = len(g)
        if n < 7:
            # name 是元组，如 ('主管A', '碳钢', '地下')
            # 构造填充行：分组列保留原值，数据列用 '/'
            padding = pd.DataFrame([
                dict(zip(group_cols, name)) | {c: '/' for c in value_cols}
                for _ in range(7 - n)
            ])
            g = pd.concat([g, padding], ignore_index=True)
        dfz.append(g)

    df_filled = pd.concat(dfz, ignore_index=True)
    df_agg = (
    df_filled.groupby(group_cols)[value_cols]
    .apply(lambda x: x.to_dict('records'))
    .rename('明细')
    .reset_index()
    )
    temp_dict={'泄漏评估报告':df_agg.to_dict(orient='records')}
    data_dict.update(temp_dict)
    return data_dict

if __name__ == "__main__":
    # 读取excel的原始数据
    CONFIG=set_argumments([
        (2,'数据源文件','xlsx',r'E:\BaiduNetdiskDownload\金牛老改\2025年金牛区评估数据.xlsx'),
        (2,'模板','docx',r'E:\BaiduNetdiskDownload\金牛老改\模板-成都-金牛-改.docx'),
        # (0,'照片文件夹','',r'E:\BaiduNetdiskDownload\金牛老改\路由图'),
        (0,'保存文件夹','',r'E:\BaiduNetdiskDownload\金牛老改\输出'),
    ])
    start = time.time()
    # df=pd.read_excel(Path(CONFIG['数据源文件']),sheet_name="Sheet5")
    # print(df.info())
    
    dfs=pd.read_excel(Path(CONFIG['数据源文件']),sheet_name=None)
    for report_no in dfs['封面']['报告编号'][:]:

        # 开启模板分析数据
        tpl = DocxTemplate(Path(CONFIG['模板']))
        data=make_data_in_list(tpl=tpl,dfs=dfs,key=report_no)
        
        # 填充并保存内容
        tpl.render(data)
        tpl.save(Path(CONFIG['保存文件夹'])/f'{report_no}.docx')
        clean_and_save(Path(CONFIG['保存文件夹'])/f'{report_no}.docx')
    
    print(f"\n总耗时: {time.time() - start:.2f} 秒")
    print('\nDone!')
